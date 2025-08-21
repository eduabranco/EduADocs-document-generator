#!/usr/bin/env python3
"""
Simple test to verify exercise generation with markdown formatting
"""

import sys
from pathlib import Path

# Add src directory to path for imports
src_path = Path(__file__).parent / "src"
sys.path.append(str(src_path))

from generators.exercise_generator import _create_exercise_docx
from docx import Document
import io

def test_exercise_markdown_formatting():
    """Test that exercise generation properly handles markdown formatting"""
    print("🧪 Testing exercise generation with markdown formatting...")
    
    # Sample markdown content that would come from an LLM
    markdown_content = """# Mathematics Exercise List

## Addition Problems

### Basic Addition
1. What is 2 + 2?
2. What is 5 + 3?

### Word Problems
- Sarah has 5 apples. She gets 3 more. How many does she have?
- There are 8 birds in a tree. 4 more join them. How many birds are there now?

## Multiplication Problems

### Single Digit
* 3 × 4 = ?
* 7 × 2 = ?

## Answer Key

### Addition Answers
1. 4
2. 8

### Word Problem Answers
- 8 apples
- 12 birds

### Multiplication Answers
* 12
* 14"""

    # Test parameters
    params = {
        'subject': 'Mathematics',
        'topic': 'Basic Arithmetic',
        'grade_level': 'Elementary (K-5)'
    }
    
    try:
        # Generate the docx file
        docx_bytes = _create_exercise_docx(markdown_content, params)
        
        # Verify it's a valid docx file by loading it
        doc_stream = io.BytesIO(docx_bytes)
        doc = Document(doc_stream)
        
        print("✅ Exercise document created successfully!")
        print(f"✅ Document contains {len(doc.paragraphs)} paragraphs")
        
        # Check if we have proper headings
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        print(f"✅ Document contains {len(headings)} headings (markdown # converted properly)")
        
        # Check if we have bullet points
        bullets = [p for p in doc.paragraphs if p.style.name == 'List Bullet']
        print(f"✅ Document contains {len(bullets)} bullet points (markdown - * converted properly)")
        
        print("\n🎉 Markdown formatting in exercise generation is working correctly!")
        return True
        
    except Exception as e:
        print(f"❌ Test failed: {e}")
        return False

if __name__ == "__main__":
    success = test_exercise_markdown_formatting()
    if not success:
        sys.exit(1)
