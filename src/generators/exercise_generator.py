from llm_handlers.api_handler import get_llm_response
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io

def generate_exercises(params):
    """Generate exercise list document"""
    
    prompt = _build_exercise_prompt(params)
    
    try:
        # Get content from LLM
        content = get_llm_response(prompt, params["llm_config"])
        
        # Create Word document
        docx_file = _create_exercise_docx(content, params)
        
        return {
            "success": True,
            "content": content,
            "docx_file": docx_file
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}

def _build_exercise_prompt(params):
    """Build prompt for exercise generation"""
    
    prompt = f"""
    Create a comprehensive exercise list for {params['subject']} at {params['grade_level']} level.
    
    Topic: {params['topic']}
    Number of questions: {params['num_questions']}
    Difficulty: {params['difficulty']}
    Question types: {', '.join(params['question_types'])}
    
    Please structure the exercises using clean Markdown formatting as follows:
    1. Start with a brief introduction to the topic
    2. Use clear heading structure:
       - # for the main title
       - ## for major sections
       - ### for subsections only
    3. Use numbered lists (1. 2. 3.) for questions
    4. Organize questions by difficulty (if mixed difficulty is selected)
    5. Include clear instructions for each section
    6. For multiple choice questions, provide 4 options (A, B, C, D)
    7. For problem-solving questions, show step-by-step solutions
    8. End with an answer key
    
    FORMATTING RULES:
    - Use # Exercise List for the main title
    - Use ## Section Name for major sections (e.g., ## Multiple Choice Questions)
    - Use ### Subsection Name only when needed
    - DO NOT use --- horizontal rules
    - DO NOT mix heading levels (like ### ## or #### ##)
    - Use 1. 2. 3. for numbered questions
    - Keep formatting simple and clean
    
    Make sure the content is age-appropriate and educationally valuable.
    """
    
    return prompt

def _create_exercise_docx(content, params):
    """Create Word document from exercise content"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading(f"{params['subject']} - Exercise List", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center alignment
    
    # Subtitle
    doc.add_heading(f"Topic: {params['topic']}", level=2)
    doc.add_heading(f"Grade Level: {params['grade_level']}", level=3)
    
    # Add content with proper markdown formatting
    _add_formatted_content_to_docx(doc, content)
    
    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io.getvalue()

def _add_formatted_content_to_docx(doc, content):
    """Add formatted content to Word document, parsing Markdown formatting"""
    lines = content.split('\n')
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        
        # Skip horizontal rules
        if line == '---' or line.startswith('---'):
            continue
        
        # Clean up malformed headings (remove extra # symbols)
        if line.startswith('#'):
            # Count the number of # at the start
            hash_count = 0
            for char in line:
                if char == '#':
                    hash_count += 1
                else:
                    break
            
            # Extract the text after all the #
            remaining_text = line[hash_count:].strip()
            
            # Remove any additional # symbols from the beginning of the text
            while remaining_text.startswith('#'):
                remaining_text = remaining_text[1:].strip()
            
            # Limit heading levels to 4
            actual_level = min(hash_count, 4)
            if actual_level > 0 and remaining_text:
                doc.add_heading(remaining_text, level=actual_level)
                continue
            
        # Check if it's a bullet point
        if line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
            doc.add_paragraph(line[2:], style='List Bullet')
        # Check if it's a numbered list
        elif len(line) > 3 and line[0].isdigit() and line[1:3] in ['. ', ') ']:
            doc.add_paragraph(line[3:], style='List Number')
        else:
            # Handle bold formatting and other text
            formatted_text = line
            
            # Remove markdown bold formatting (**text**)
            import re
            formatted_text = re.sub(r'\*\*(.*?)\*\*', r'\1', formatted_text)
            
            doc.add_paragraph(formatted_text)