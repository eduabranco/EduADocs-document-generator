from llm_handlers.api_handler import get_llm_response
from docx import Document
import io

def generate_summary(params):
    """Generate summary document"""
    
    prompt = _build_summary_prompt(params)
    
    try:
        # Get content from LLM
        content = get_llm_response(prompt, params["llm_config"])
        
        # Create Word document
        docx_file = _create_summary_docx(content, params)
        
        return {
            "success": True,
            "content": content,
            "docx_file": docx_file
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}

def _build_summary_prompt(params):
    """Build prompt for summary generation"""
    
    prompt = f"""
    Create a comprehensive summary for {params['subject']} at {params['grade_level']} level.
    
    Topic: {params['topic']}
    Length: {params['summary_length']}
    Format: {params['format_style']}
    Include examples: {params['include_examples']}
    
    Structure the summary with:
    1. Introduction to the topic
    2. Main concepts and key points
    3. {'Real-world examples and applications' if params['include_examples'] else 'Theoretical explanations'}
    4. Summary of key takeaways
    5. Suggested further reading or activities
    
    Make sure the content is:
    - Age-appropriate for {params['grade_level']}
    - Well-organized and easy to follow
    - Educationally comprehensive
    - Formatted according to {params['format_style']} style
    """
    
    return prompt

def _create_summary_docx(content, params):
    """Create Word document from summary content"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading(f"{params['subject']} - Summary", 0)
    title.alignment = 1  # Center alignment
    
    # Subtitle
    doc.add_heading(f"Topic: {params['topic']}", level=2)
    doc.add_heading(f"Grade Level: {params['grade_level']}", level=3)
    
    # Add content based on format style
    if params['format_style'] == "Bullet Points":
        _add_bullet_content(doc, content)
    elif params['format_style'] == "Outline":
        _add_outline_content(doc, content)
    elif params['format_style'] == "Q&A Format":
        _add_qa_content(doc, content)
    else:  # Paragraphs
        _add_paragraph_content(doc, content)
    
    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io.getvalue()

def _add_bullet_content(doc, content):
    """Add content in bullet point format"""
    _add_formatted_content_to_docx(doc, content)

def _add_outline_content(doc, content):
    """Add content in outline format"""
    _add_formatted_content_to_docx(doc, content)

def _add_qa_content(doc, content):
    """Add content in Q&A format"""
    _add_formatted_content_to_docx(doc, content)

def _add_paragraph_content(doc, content):
    """Add content in paragraph format"""
    _add_formatted_content_to_docx(doc, content)

def _add_formatted_content_to_docx(doc, content):
    """Add formatted content to Word document, parsing Markdown formatting"""
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
            
        # Check if it's a heading (starts with #)
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        # Check if it's a bullet point
        elif line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
            doc.add_paragraph(line[2:], style='List Bullet')
        # Check if it's a numbered list
        elif line[0].isdigit() and line[1:3] in ['. ', ') ']:
            doc.add_paragraph(line[3:], style='List Number')
        else:
            doc.add_paragraph(line)